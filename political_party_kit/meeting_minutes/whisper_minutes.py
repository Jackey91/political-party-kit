"""Meeting minutes generation powered by Whisper and GPT.

This module provides utilities to transcribe meeting audio recordings with
OpenAI's Whisper API, summarise the results with GPT models and export a
structured protocol as a Word document.  The functionality is based on an
internal script that has been promoted to a reusable module so that other
parts of the Political Party Kit can integrate with it.
"""

from __future__ import annotations

import copy
import os
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openai import OpenAI

try:  # pragma: no cover - optional dependency
    import tiktoken  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    tiktoken = None  # type: ignore

__all__ = [
    "MeetingMetadata",
    "MeetingMinutesGenerator",
    "collect_meeting_metadata",
    "parse_semicolon_list",
]


@dataclass
class MeetingMetadata:
    """Metadata that is attached to the generated meeting minutes."""

    titel: str = "Sitzungsprotokoll"
    datum: Optional[str] = None
    ort: Optional[str] = None
    startzeit: Optional[str] = None
    endzeit: Optional[str] = None
    moderation: Optional[str] = None
    protokoll: Optional[str] = None
    teilnehmer: List[str] = field(default_factory=list)
    gaeste: List[str] = field(default_factory=list)
    agenda: List[str] = field(default_factory=list)
    ziele: List[str] = field(default_factory=list)
    hinweise: List[str] = field(default_factory=list)

    def to_prompt_header(self) -> str:
        """Return a multi-line string describing the session for GPT prompts."""

        lines: List[str] = []
        if self.titel:
            lines.append(f"Titel: {self.titel}")
        if self.datum:
            lines.append(f"Datum: {self.datum}")
        slot = " - ".join(filter(None, [self.startzeit, self.endzeit]))
        if slot:
            lines.append(f"Zeitfenster: {slot}")
        if self.ort:
            lines.append(f"Ort: {self.ort}")
        if self.moderation:
            lines.append(f"Moderation: {self.moderation}")
        if self.protokoll:
            lines.append(f"Protokollführung: {self.protokoll}")
        if self.teilnehmer:
            lines.append(f"Teilnehmer: {human_join(self.teilnehmer)}")
        if self.gaeste:
            lines.append(f"Gäste: {human_join(self.gaeste)}")
        if self.ziele:
            lines.append("Ziele:")
            lines.extend([f"  - {goal}" for goal in self.ziele])
        if self.agenda:
            lines.append("Tagesordnung:")
            lines.extend([f"  - {item}" for item in self.agenda])
        if self.hinweise:
            lines.append("Hinweise:")
            lines.extend([f"  - {note}" for note in self.hinweise])
        return "\n".join(lines)

    def to_docx_header_fields(self) -> List[str]:
        """Return the metadata elements that should be shown at the top of the document."""

        entries: List[str] = []
        if self.datum:
            entries.append(f"Datum: {self.datum}")
        slot = " - ".join(filter(None, [self.startzeit, self.endzeit]))
        if slot:
            entries.append(f"Zeitfenster: {slot}")
        if self.ort:
            entries.append(f"Ort: {self.ort}")
        if self.moderation:
            entries.append(f"Moderation: {self.moderation}")
        if self.protokoll:
            entries.append(f"Protokoll: {self.protokoll}")
        if self.teilnehmer:
            entries.append(f"Teilnehmer: {human_join(self.teilnehmer)}")
        if self.gaeste:
            entries.append(f"Gäste: {human_join(self.gaeste)}")
        return entries

    def to_dict(self) -> Dict[str, Any]:
        """Return a dictionary representation compatible with the legacy script."""

        return {
            "titel": self.titel,
            "datum": self.datum,
            "ort": self.ort,
            "startzeit": self.startzeit,
            "endzeit": self.endzeit,
            "moderation": self.moderation,
            "protokoll": self.protokoll,
            "teilnehmer": list(self.teilnehmer),
            "gaeste": list(self.gaeste),
            "agenda": list(self.agenda),
            "ziele": list(self.ziele),
            "hinweise": list(self.hinweise),
        }

    @classmethod
    def from_dict(cls, payload: Dict[str, Any]) -> "MeetingMetadata":
        """Create metadata from a dictionary."""

        return cls(
            titel=payload.get("titel", "Sitzungsprotokoll"),
            datum=payload.get("datum") or None,
            ort=payload.get("ort") or None,
            startzeit=payload.get("startzeit") or None,
            endzeit=payload.get("endzeit") or None,
            moderation=payload.get("moderation") or None,
            protokoll=payload.get("protokoll") or None,
            teilnehmer=list(payload.get("teilnehmer", []) or []),
            gaeste=list(payload.get("gaeste", []) or []),
            agenda=list(payload.get("agenda", []) or []),
            ziele=list(payload.get("ziele", []) or []),
            hinweise=list(payload.get("hinweise", []) or []),
        )


def human_join(items: Iterable[str]) -> str:
    """Join human readable names."""

    clean = [item.strip() for item in items if item and item.strip()]
    return ", ".join(clean)


def parse_semicolon_list(value: str) -> List[str]:
    """Convert a semicolon-separated list to a list of trimmed strings."""

    if not value:
        return []
    return [item.strip() for item in value.split(";") if item.strip()]


def load_api_key() -> str:
    """Load the OpenAI API key from the environment or a .env file."""

    load_dotenv()
    key = os.getenv("OPENAI_API_KEY", "").strip()
    if not key:
        print(
            "Fehler: Kein OPENAI_API_KEY gefunden. Bitte .env erstellen oder Umgebungsvariable setzen.",
            file=sys.stderr,
        )
        raise SystemExit(1)
    return key


def chunk_text(text: str, target_chars: int = 7000) -> List[str]:
    """Split text into chunks respecting paragraph boundaries."""

    if not text:
        return []
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks: List[str] = []
    buffer = ""
    for paragraph in paragraphs:
        candidate = f"{buffer}\n{paragraph}".strip() if buffer else paragraph
        if len(candidate) <= target_chars:
            buffer = candidate
        else:
            if buffer:
                chunks.append(buffer)
            buffer = paragraph
    if buffer:
        chunks.append(buffer)
    return chunks


def transcribe_audio(client: OpenAI, audio_path: str, language_hint: str = "de") -> str:
    """Transcribe an audio file using Whisper."""

    file_path = Path(audio_path)
    if not file_path.exists():
        raise FileNotFoundError(f"Audiodatei nicht gefunden: {audio_path}")

    with file_path.open("rb") as handle:
        transcription = client.audio.transcriptions.create(
            model="whisper-1",
            file=handle,
            language=language_hint,
            response_format="text",
        )
    if isinstance(transcription, str):
        return transcription
    return getattr(transcription, "text", str(transcription))


def summarize_chunk(client: OpenAI, text: str) -> str:
    """Summarise a text chunk using GPT."""

    system_prompt = (
        "Du bist ein Assistent, der deutsche Sitzungsprotokolle erstellt. "
        "Fasse den folgenden Text klar, präzise und sachlich zusammen. "
        "Nenne relevante Themen, Argumente, Entscheidungen und Aufgaben mit Zuständigkeiten, wenn erkennbar."
    )
    user_prompt = f"Textauszug:\n\n{text}"
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.2,
    )
    return response.choices[0].message.content.strip()


def consolidate_summaries(client: OpenAI, summaries: List[str], meta: MeetingMetadata) -> str:
    """Combine individual summaries into a final protocol."""

    system_prompt = (
        "Erstelle ein vollständiges, gut strukturiertes deutsches Sitzungsprotokoll aus den Teilzusammenfassungen. "
        "Gliedere in: 1) Teilnehmer 2) Agenda/Themenblöcke 3) Diskussion (kurz) 4) Beschlüsse 5) Aufgaben & Zuständigkeiten 6) Nächste Schritte. "
        "Arbeite stichpunktorientiert, aber vollständig; vermeide irrelevante Details. "
        "Nutze neutrale, präzise Sprache und konsistente Formatierung."
    )
    user_prompt = f"{meta.to_prompt_header()}\n\nTeilzusammenfassungen:\n\n{'\n\n'.join(summaries)}"
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.2,
    )
    return response.choices[0].message.content.strip()


def build_docx(text: str, out_path: str, meta: MeetingMetadata) -> None:
    """Create a Word document with the final minutes."""

    document = Document()

    title = meta.titel or "Sitzungsprotokoll"
    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_entries = meta.to_docx_header_fields()
    if sub_entries:
        sub_paragraph = document.add_paragraph("  |  ".join(sub_entries))
        sub_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")

    def add_section_heading(value: str) -> None:
        heading = document.add_paragraph()
        run = heading.add_run(value)
        run.bold = True
        run.font.size = Pt(12)

    if meta.agenda:
        add_section_heading("Tagesordnung")
        for entry in meta.agenda:
            document.add_paragraph(entry, style="List Number")
        document.add_paragraph("")

    if meta.ziele:
        add_section_heading("Sitzungsziele")
        for goal in meta.ziele:
            document.add_paragraph(goal, style="List Bullet")
        document.add_paragraph("")

    if meta.hinweise:
        add_section_heading("Besondere Hinweise")
        for note in meta.hinweise:
            document.add_paragraph(note, style="List Bullet")
        document.add_paragraph("")

    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith(("#", "##", "###")):
            level = stripped.count("#")
            paragraph = document.add_paragraph()
            run = paragraph.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.size = Pt(14 if level == 1 else 12)
        elif stripped.endswith(":") and len(stripped) < 120:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(stripped)
            run.bold = True
        else:
            if stripped.startswith(("-", "*", "•")):
                document.add_paragraph(stripped.lstrip("-*• ").strip(), style="List Bullet")
            else:
                document.add_paragraph(line)

    document.add_paragraph("")
    stamp = document.add_paragraph(f"Erstellt am {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    stamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.save(out_path)


def prompt_with_default(prompt: str, default: str = "") -> str:
    """Prompt the user for a value with a default."""

    hint = f" [{default}]" if default else ""
    value = input(f"{prompt}{hint}: ").strip()
    return value or default


def prompt_list(prompt: str, default: List[str]) -> List[str]:
    """Prompt the user for a list of values."""

    print(f"{prompt}:")
    if default:
        print("  (Aktuell hinterlegt:)")
        for idx, item in enumerate(default, start=1):
            print(f"    {idx}. {item}")
    print("  Geben Sie pro Zeile einen neuen Eintrag ein. Leer lassen, um Eingabe zu beenden und ggf. bestehende Werte zu übernehmen.")
    entries: List[str] = []
    counter = 1
    while True:
        value = input(f"    {counter}. ").strip()
        if not value:
            break
        entries.append(value)
        counter += 1
    return entries if entries else default


def collect_meeting_metadata(base_meta: MeetingMetadata) -> MeetingMetadata:
    """Interactively collect meeting metadata from the user."""

    meta = copy.deepcopy(base_meta)
    while True:
        print("\n=== Eingabemaske für Sitzungsbeginn ===")
        meta.titel = prompt_with_default("Sitzungstitel", meta.titel or "Sitzungsprotokoll")
        meta.datum = prompt_with_default(
            "Datum (z. B. 2025-10-01)", meta.datum or datetime.now().strftime("%Y-%m-%d")
        )
        meta.startzeit = prompt_with_default("Geplanter Beginn (HH:MM)", meta.startzeit or "")
        meta.endzeit = prompt_with_default("Geplantes Ende (HH:MM)", meta.endzeit or "")
        meta.ort = prompt_with_default("Ort", meta.ort or "")
        meta.moderation = prompt_with_default("Leitung/Moderation", meta.moderation or "")
        meta.protokoll = prompt_with_default("Protokollführung", meta.protokoll or "")
        meta.teilnehmer = prompt_list("Teilnehmer*innen", meta.teilnehmer)
        meta.gaeste = prompt_list("Gäste", meta.gaeste)
        meta.agenda = prompt_list("Tagesordnungspunkte", meta.agenda)
        meta.ziele = prompt_list("Sitzungsziele", meta.ziele)
        meta.hinweise = prompt_list("Besondere Hinweise", meta.hinweise)

        print("\nZusammenfassung der Angaben:")
        summary_lines = [
            f"  Titel: {meta.titel or ''}",
            f"  Datum: {meta.datum or ''}",
            f"  Zeitfenster: {' - '.join(filter(None, [meta.startzeit, meta.endzeit]))}",
            f"  Ort: {meta.ort or ''}",
            f"  Moderation: {meta.moderation or ''}",
            f"  Protokollführung: {meta.protokoll or ''}",
            f"  Teilnehmer*innen: {human_join(meta.teilnehmer)}",
            f"  Gäste: {human_join(meta.gaeste)}",
            f"  Ziele: {human_join(meta.ziele)}",
            f"  Hinweise: {human_join(meta.hinweise)}",
        ]
        for line in summary_lines:
            print(line)
        if meta.agenda:
            print("  Tagesordnung:")
            for idx, item in enumerate(meta.agenda, start=1):
                print(f"    {idx}. {item}")

        confirm = input("\nSind die Angaben korrekt? (J/n): ").strip().lower()
        if confirm in ("", "j", "ja", "y", "yes"):
            return meta
        print("\nBitte Eingaben erneut vornehmen.\n")


class MeetingMinutesGenerator:
    """High level orchestrator that wires the helper functions together."""

    def __init__(self, client: Optional[OpenAI] = None) -> None:
        self.client = client or OpenAI(api_key=load_api_key())

    def generate(
        self,
        audio_path: str,
        output_path: str,
        metadata: MeetingMetadata,
        language_hint: str = "de",
        save_transcript: Optional[str] = None,
        save_partials: Optional[str] = None,
        prompt_for_metadata: bool = False,
    ) -> str:
        """Generate meeting minutes from an audio file.

        Returns the path to the generated document.
        """

        meta = copy.deepcopy(metadata)
        if prompt_for_metadata:
            meta = collect_meeting_metadata(meta)

        transcript = transcribe_audio(self.client, audio_path, language_hint=language_hint).strip()

        if save_transcript:
            Path(save_transcript).write_text(transcript, encoding="utf-8")

        chunks = chunk_text(transcript, target_chars=7000)

        partial_summaries: List[str] = []
        if chunks:
            partial_dir = Path(save_partials) if save_partials else None
            if partial_dir and not partial_dir.exists():
                partial_dir.mkdir(parents=True, exist_ok=True)

            for index, chunk in enumerate(chunks, start=1):
                summary = summarize_chunk(self.client, chunk)
                partial_summaries.append(summary)
                if partial_dir:
                    (partial_dir / f"teil_{index:02d}.md").write_text(summary, encoding="utf-8")

        if partial_summaries:
            final_minutes = consolidate_summaries(self.client, partial_summaries, meta)
        else:
            final_minutes = consolidate_summaries(self.client, [transcript], meta)

        build_docx(final_minutes, output_path, meta)
        return output_path


__all__.extend([name for name in globals() if name.endswith("MinutesGenerator")])
